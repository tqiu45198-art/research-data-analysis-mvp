import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor

def create_research_analysis_bp():
    # 1. 初始化文档与页面设置（A4尺寸）
    doc = Document()
    section = doc.sections[0]
    section.page_width = Pt(595.3)  # A4宽度
    section.page_height = Pt(841.9) # A4高度
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)

    # 2. 封面设计
    # 封面标题
    cover_heading = doc.add_heading("科研数据智能分析平台", 0)
    cover_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cover_run = cover_heading.runs[0]
    cover_run.font.size = Pt(32)
    cover_run.font.color.rgb = RGBColor(30, 144, 255)  # 科技蓝
    cover_run.font.name = "微软雅黑"
    cover_run.element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 封面副标题
    sub_heading = doc.add_paragraph("——低代码·大模型驱动的本科生科研赋能工具")
    sub_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_heading.runs[0]
    sub_run.font.size = Pt(16)
    sub_run.font.name = "微软雅黑"
    sub_run.element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 封面信息（团队/领域）
    doc.add_paragraph("\n\n\n\n\n\n\n\n")  # 换行占位
    team_info = doc.add_paragraph("团队：上海工程技术大学数理与统计学院本科生团队")
    team_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    field_info = doc.add_paragraph("领域：人工智能赋能教育（科技创新与未来产业）")
    field_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in team_info.runs + field_info.runs:
        run.font.size = Pt(12)
        run.font.name = "微软雅黑"
        run.element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 3. 目录页
    doc.add_page_break()
    toc_heading = doc.add_heading("目录", 1)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_heading.runs[0].font.name = "微软雅黑"
    toc_heading.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 目录内容（对应后续章节）
    toc_items = [
        "1. 执行摘要",
        "2. 项目背景与问题识别",
        "3. 产品与解决方案",
        "4. 技术方案与创新点",
        "5. 市场分析与发展战略",
        "6. 团队介绍",
        "7. 融资需求与财务规划",
        "8. 风险与应对措施",
        "9. 发展规划与里程碑"
    ]
    for item in toc_items:
        para = doc.add_paragraph(item)
        para.runs[0].font.size = Pt(11)
        para.runs[0].font.name = "微软雅黑"
        para.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 4. 执行摘要
    doc.add_page_break()
    summary_heading = doc.add_heading("1. 执行摘要", 1)
    summary_heading.runs[0].font.name = "微软雅黑"
    summary_heading.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    summary_content = """
    项目领域：科技创新和未来产业（人工智能赋能教育）

    团队构成：以上海工程技术大学数理与统计学院学生为核心，由11名数据计算与应用专业本科生组成跨职能团队，由学院讲师指导。

    核心痛点：本科生科研存在显著数据分析技术鸿沟——传统统计软件（SPSS、Stata）操作复杂，Python/R等编程语言门槛过高，阻碍无编程基础的文理科学生开展数据驱动研究；即使获得计算结果，学生常难以理解统计学含义，无法转化为研究结论。

    解决方案：开发低代码、大模型驱动的科研数据智能分析平台，融合低代码交互、大语言模型自然语言理解、专业统计计算库，提供“数据上传→需求描述→智能分析→报告生成”全流程服务。用户通过自然语言（如“分析两种教学方法对成绩的差异”）即可获得专业解读，极大降低技术门槛。

    市场前景：核心市场为全国高校本科生（潜在用户超1000万），需求刚性且供给匮乏。规划“校内试点→区域推广→全国拓展”路径，致力于成为科研创新基础工具。
    """
    summary_para = doc.add_paragraph(summary_content)
    summary_para.runs[0].font.size = Pt(11)
    summary_para.runs[0].font.name = "微软雅黑"
    summary_para.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    summary_para.space_after = Pt(6)

    # 5. 项目背景与问题识别
    doc.add_page_break()
    bg_heading = doc.add_heading("2. 项目背景与问题识别", 1)
    bg_heading.runs[0].font.name = "微软雅黑"
    bg_heading.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 5.1 背景
    bg_sub1 = doc.add_heading("2.1 项目背景", 2)
    bg_sub1.runs[0].font.name = "微软雅黑"
    bg_sub1.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    bg_content1 = """
    在跨学科创新与数据驱动研究的背景下，本科生参与科研成为创新能力培养的关键环节。然而，多数本科生面临“研究设想→数据结论”的转化障碍，核心问题集中于“工具使用”与“知识理解”两大层面。
    """
    bg_para1 = doc.add_paragraph(bg_content1)
    bg_para1.runs[0].font.size = Pt(11)
    bg_para1.runs[0].font.name = "微软雅黑"
    bg_para1.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 5.2 核心痛点
    bg_sub2 = doc.add_heading("2.2 核心痛点", 2)
    bg_sub2.runs[0].font.name = "微软雅黑"
    bg_sub2.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    pain_points = [
        "工具使用障碍：SPSS、Stata菜单繁杂，Python/R需大量编程学习，学生精力从研究设计转移至技术实现；",
        "知识理解断层：即使获得计算结果，学生难以理解统计学含义（如p值、R²），无法转化为有意义的研究结论；",
        "能力真空地带：“研究问题”与“数据答案”间缺乏桥梁，数据分析成为本科生科研的核心壁垒。"
    ]
    for point in pain_points:
        para = doc.add_paragraph(f"• {point}")
        para.runs[0].font.size = Pt(11)
        para.runs[0].font.name = "微软雅黑"
        para.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 6. 产品与解决方案
    doc.add_page_break()
    product_heading = doc.add_heading("3. 产品与解决方案", 1)
    product_heading.runs[0].font.name = "微软雅黑"
    product_heading.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 6.1 核心功能
    product_sub1 = doc.add_heading("3.1 核心功能", 2)
    product_sub1.runs[0].font.name = "微软雅黑"
    product_sub1.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    functions = [
        "低代码可视化操作：拖拽式数据上传、勾选式分析设置，无需代码即可完成数据预处理与分析配置；",
        "自然语言交互：支持日常语言输入需求（如“分析城市对订单量的影响”），自动匹配t检验、方差分析等方法；",
        "自动化统计引擎：集成Scikit-learn、Statsmodels等库，覆盖描述性统计、回归、聚类等10+种分析方法；",
        "结构化智能报告：输出含图表、原理解读、显著性判断的报告，而非杂乱数据（如“p<0.05代表差异显著”）。"
    ]
    for func in functions:
        para = doc.add_paragraph(f"• {func}")
        para.runs[0].font.size = Pt(11)
        para.runs[0].font.name = "微软雅黑"
        para.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 6.2 服务模式
    product_sub2 = doc.add_heading("3.2 服务模式", 2)
    product_sub2.runs[0].font.name = "微软雅黑"
    product_sub2.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    service_mode = """
    • 免费基础版：满足课程作业、小组项目需求（如描述性统计、简单t检验），吸引海量用户构建社区；
    • 付费专业版（SaaS订阅）：面向大创、挑战杯等深度科研项目，提供高级模型（如多项式回归、3D聚类）、私有化部署、优先技术支持，按年订阅收费（预计199元/年/用户）。
    """
    service_para = doc.add_paragraph(service_mode)
    service_para.runs[0].font.size = Pt(11)
    service_para.runs[0].font.name = "微软雅黑"
    service_para.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 7. 技术方案与创新点
    doc.add_page_break()
    tech_heading = doc.add_heading("4. 技术方案与创新点", 1)
    tech_heading.runs[0].font.name = "微软雅黑"
    tech_heading.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 7.1 系统架构
    tech_sub1 = doc.add_heading("4.1 三层解耦架构", 2)
    tech_sub1.runs[0].font.name = "微软雅黑"
    tech_sub1.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    architecture = [
        "交互层：基于Streamlit框架构建响应式Web界面，实现低代码拖拽、需求输入等交互逻辑；",
        "认知与调度层：以LangChain为核心，集成LoRA微调的Llama 3-8B模型，专注科研场景需求理解（如识别“差异分析”对应t检验）；",
        "计算层：封装Pandas、NumPy、Scikit-learn等库，确保统计计算的准确性与高性能，支持并行处理多文件数据。"
    ]
    for arch in architecture:
        para = doc.add_paragraph(f"• {arch}")
        para.runs[0].font.size = Pt(11)
        para.runs[0].font.name = "微软雅黑"
        para.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 7.2 创新点
    tech_sub2 = doc.add_heading("4.2 三大创新点", 2)
    tech_sub2.runs[0].font.name = "微软雅黑"
    tech_sub2.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    innovations = [
        "模式创新：首创“低代码前端+大模型中台+专业计算后端”模式，平衡易用性、智能性与专业性；",
        "场景创新：聚焦“本科生科研”细分市场，产品设计贴合无编程基础用户需求（如自然语言交互、结果解读）；",
        "价值创新：不仅提供分析结果，更通过报告解读提升用户数据素养，充当“随身数据分析导师”。"
    ]
    for inn in innovations:
        para = doc.add_paragraph(f"• {inn}")
        para.runs[0].font.size = Pt(11)
        para.runs[0].font.name = "微软雅黑"
        para.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 8. 市场分析与发展战略
    doc.add_page_break()
    market_heading = doc.add_heading("5. 市场分析与发展战略", 1)
    market_heading.runs[0].font.name = "微软雅黑"
    market_heading.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 8.1 目标市场
    market_sub1 = doc.add_heading("5.1 目标市场", 2)
    market_sub1.runs[0].font.name = "微软雅黑"
    market_sub1.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    market_content = """
    核心市场：全国高校本科生（潜在用户规模超1000万），重点覆盖文理科学生（如经济学、生物学、社会学等需数据分析的专业）；
    拓展市场：中期向企业初级研发、市场调研场景延伸（如中小企业轻量级数据分析需求）。
    """
    market_para = doc.add_paragraph(market_content)
    market_para.runs[0].font.size = Pt(11)
    market_para.runs[0].font.name = "微软雅黑"
    market_para.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 8.2 竞争分析
    market_sub2 = doc.add_heading("5.2 竞争分析", 2)
    market_sub2.runs[0].font.name = "微软雅黑"
    market_sub2.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    competition = [
        "vs 传统软件（SPSS/Stata）：更智能（自然语言交互）、更易用（零代码）、交互更贴合年轻人习惯；",
        "vs 编程语言（Python/R）：零门槛，用户无需学习语法，专注研究逻辑；",
        "vs 通用AI（ChatGPT）：科研领域更专业（结果可验证、统计方法匹配准确），注重数据隐私（本地/私有部署）。"
    ]
    for comp in competition:
        para = doc.add_paragraph(f"• {comp}")
        para.runs[0].font.size = Pt(11)
        para.runs[0].font.name = "微软雅黑"
        para.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 8.3 发展战略
    market_sub3 = doc.add_heading("5.3 三阶段发展战略", 2)
    market_sub3.runs[0].font.name = "微软雅黑"
    market_sub3.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    strategy = [
        "第一阶段（0-6个月）：校内试点，嵌入数理学院《数据分析》课程，积累500+种子用户，完成2轮产品迭代；",
        "第二阶段（6-18个月）：以上海高校为突破口，通过学术竞赛（挑战杯、大创）、培训工作坊推广，覆盖长三角50+高校；",
        "第三阶段（18-36个月）：全国拓展，与200+高校达成合作，推出企业版，实现年营收超500万元。"
    ]
    for strat in strategy:
        para = doc.add_paragraph(f"• {strat}")
        para.runs[0].font.size = Pt(11)
        para.runs[0].font.name = "微软雅黑"
        para.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 9. 团队介绍
    doc.add_page_break()
    team_heading = doc.add_heading("6. 团队介绍", 1)
    team_heading.runs[0].font.name = "微软雅黑"
    team_heading.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    team_structure = [
        "项目负责人：1名，负责战略规划、进度管理与资源协调；",
        "技术研发组（6人）：分模型算法、统计计算、前端开发3方向，成员具备计算数学、统计学专业背景；",
        "需求与数据组（3人）：负责市场调研、用户需求分析、测试案例构建；",
        "测试与推广组（2人）：负责产品测试、用户体验优化、校园推广活动；",
        "指导教师：1名学院讲师，提供统计学理论指导与学术资源支持。"
    ]
    for role in team_structure:
        para = doc.add_paragraph(f"• {role}")
        para.runs[0].font.size = Pt(11)
        para.runs[0].font.name = "微软雅黑"
        para.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 10. 融资需求与财务规划
    doc.add_page_break()
    finance_heading = doc.add_heading("7. 融资需求与财务规划", 1)
    finance_heading.runs[0].font.name = "微软雅黑"
    finance_heading.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 10.1 融资需求
    finance_sub1 = doc.add_heading("7.1 融资需求（种子轮）", 2)
    finance_sub1.runs[0].font.name = "微软雅黑"
    finance_sub1.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    finance_demand = """
    计划融资200万元，资金使用分配如下：
    • 产品研发与技术支持（40%）：云服务器租赁、第三方API服务、UI/UX优化、模型微调；
    • 市场推广与用户获取（30%）：校园活动、线上内容营销（知乎/小红书）、渠道合作；
    • 团队建设与运营（20%）：核心成员津贴、团队建设、办公耗材；
    • 风险储备金（10%）：应对未预见开支（如服务器扩容、紧急迭代）。
    """
    finance_para1 = doc.add_paragraph(finance_demand)
    finance_para1.runs[0].font.size = Pt(11)
    finance_para1.runs[0].font.name = "微软雅黑"
    finance_para1.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 10.2 收入预测
    finance_sub2 = doc.add_heading("7.2 收入预测（未来3年）", 2)
    finance_sub2.runs[0].font.name = "微软雅黑"
    finance_sub2.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    revenue = [
        "第1年：免费获客为主，付费用户5000+，营收50万元（主要来自高校定制合作）；",
        "第2年：付费用户5万+，营收200万元（订阅费199元/年/用户）；",
        "第3年：付费用户20万+，营收800万元（含企业版订阅收入）。"
    ]
    for rev in revenue:
        para = doc.add_paragraph(f"• {rev}")
        para.runs[0].font.size = Pt(11)
        para.runs[0].font.name = "微软雅黑"
        para.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 11. 风险与应对措施
    doc.add_page_break()
    risk_heading = doc.add_heading("8. 风险与应对措施", 1)
    risk_heading.runs[0].font.name = "微软雅黑"
    risk_heading.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    risks = [
        "技术风险：大模型需求理解准确率不足→应对：持续用科研场景数据微调模型，构建1000+标注需求库；",
        "市场风险：高校合作推进缓慢→应对：先从学生社团、竞赛切入，以“免费试用”打开合作缺口；",
        "竞争风险：大厂推出同类产品→应对：聚焦本科生细分场景，深耕“教育+科研”垂直需求，建立用户粘性。"
    ]
    for risk in risks:
        para = doc.add_paragraph(f"• {risk}")
        para.runs[0].font.size = Pt(11)
        para.runs[0].font.name = "微软雅黑"
        para.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 12. 发展规划与里程碑
    doc.add_page_break()
    plan_heading = doc.add_heading("9. 发展规划与里程碑", 1)
    plan_heading.runs[0].font.name = "微软雅黑"
    plan_heading.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    milestones = [
        " Month 3：完成V1.0版本开发，支持5种核心分析方法，校内试点启动；",
        " Month 6：用户数突破1000，完成V2.0版本（新增自然语言交互）；",
        " Month 12：覆盖上海20+高校，付费用户突破1万，启动长三角推广；",
        " Month 24：全国覆盖100+高校，推出企业版，营收突破200万元；",
        " Month 36：成为本科生科研数据分析头部工具，市场份额超30%。"
    ]
    for mile in milestones:
        para = doc.add_paragraph(f"• {mile}")
        para.runs[0].font.size = Pt(11)
        para.runs[0].font.name = "微软雅黑"
        para.runs[0].element.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")

    # 13. 保存文档（桌面路径）
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    if not os.path.exists(desktop_path):
        os.makedirs(desktop_path)
    file_path = os.path.join(desktop_path, "科研数据智能分析平台商业计划书.docx")
    doc.save(file_path)
    print(f"✅ 商业计划书已生成，路径：{file_path}")

if __name__ == "__main__":
    create_research_analysis_bp()
